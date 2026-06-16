#!/usr/bin/env python3
"""Generate a concise DOCX report for lab exercise 6."""

from __future__ import annotations

import struct
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "sprawozdanie_cwiczenie_6.docx"

EXERCISE_COMMANDS = [
    (
        "Na podstawie zamieszczonego w instrukcji wstępnego diagramu ERD dla "
        "systemu obsługi zamówień zaproponować przykładowy diagram hierarchii "
        "funkcji (FHD) dla minimum 5 głównych podfunkcji. Do stworzenia diagramów "
        "wykorzystać narzędzie UMLet."
    ),
    (
        "Dokonać dekompozycji funkcji głównych odpowiadających za encje: "
        "Magazyn, Rejon, Klient."
    ),
    (
        "Na podstawie stworzonych diagramów hierarchii funkcji stworzyć diagramy "
        "przepływu danych z pomocą Data Modeler: diagram kontekstowy dla systemu "
        "Obsługi zamówień, diagram systemowy dla wszystkich głównych funkcji oraz "
        "trzy diagramy szczegółowe dla wybranych procesów głównych."
    ),
    (
        "Przygotować raport zawierający tytuł ćwiczenia, autora, treść zadania, "
        "diagramy FHD oraz diagramy DFD."
    ),
]

DIAGRAMS = [
    ("Diagram FHD - funkcje główne", "fhd_01_glowne_funkcje.png"),
    (
        "Diagram FHD - dekompozycja funkcji Klient, Rejon i Magazyn",
        "fhd_02_dekompozycja_magazyn_rejon_klient.png",
    ),
    ("DFD - diagram kontekstowy", "dfd_01_kontekstowy.png"),
    ("DFD - diagram systemowy", "dfd_02_systemowy.png"),
    ("DFD szczegółowy - Obsługa klientów", "dfd_03_szczegolowy_klient.png"),
    ("DFD szczegółowy - Obsługa rejonów", "dfd_04_szczegolowy_rejon.png"),
    ("DFD szczegółowy - Obsługa magazynu", "dfd_05_szczegolowy_magazyn.png"),
]


def png_dimensions(path: Path) -> tuple[int, int]:
    data = path.read_bytes()
    if data[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError(f"Not a PNG file: {path}")
    return struct.unpack(">II", data[16:24])


def image_width_for_page(path: Path, max_width_cm: float = 16.5, max_height_cm: float = 20.0) -> float:
    width_px, height_px = png_dimensions(path)
    aspect = width_px / height_px
    return min(max_width_cm, max_height_cm * aspect)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "Ćwiczenie 6\nDiagramy hierarchii funkcji i diagramy przepływu danych"
    )
    run.bold = True
    run.font.size = Pt(17)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Skromne sprawozdanie - system obsługi zamówień").italic = True

    doc.add_paragraph("Autor: ................................................")


def add_commands(doc: Document) -> None:
    doc.add_heading("Polecenia do ćwiczenia", level=1)
    for command in EXERCISE_COMMANDS:
        doc.add_paragraph(command, style="List Number")


def add_diagrams(doc: Document) -> None:
    doc.add_heading("Diagramy", level=1)
    for index, (title, filename) in enumerate(DIAGRAMS, start=1):
        doc.add_heading(f"{index}. {title}", level=2)

        image_path = BASE_DIR / filename
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        width_cm = image_width_for_page(image_path)
        paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))

        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"Rys. {index}. {title}.")
        run.italic = True
        run.font.size = Pt(9)

        if index != len(DIAGRAMS):
            doc.add_page_break()


def main() -> None:
    doc = Document()
    set_document_defaults(doc)
    add_title(doc)
    add_commands(doc)
    doc.add_page_break()
    add_diagrams(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
